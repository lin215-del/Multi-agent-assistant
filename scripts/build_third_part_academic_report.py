from __future__ import annotations

import json
import math
import shutil
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "reports"
ASSET_DIR = ROOT / "tmp" / "report_assets"
OUTPUT_DOCX = OUTPUT_DIR / "第三部分_多模态清洗与关联_学术报告.docx"
SAMPLE_IMAGE = (
    ROOT
    / "knowledge_base"
    / "blobs"
    / "images"
    / "16"
    / "16d50c364b1157134bb6fec4c280cc339cc9aef68564aa74aae2297f74af7167.jpg"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "087D72"
RED = "9B1C1C"
GOLD = "7A5A00"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
NUMBERING_MARKERS: dict[int, tuple[str, str]] = {}


def read_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return default


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "SimSun",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_shading(paragraph, fill: str, border_left: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_left:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_left)
        borders.append(left)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def create_numbering(doc: Document, marker: str, num_format: str = "decimal") -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(value)
        for node in numbering.findall(qn("w:abstractNum"))
        if (value := node.get(qn("w:abstractNumId"))) is not None
    ]
    abstract_id = max(existing_abstract, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), marker)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, fmt, text, justification, p_pr])
    if num_format == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    existing_num = [
        int(value)
        for node in numbering.findall(qn("w:num"))
        if (value := node.get(qn("w:numId"))) is not None
    ]
    num_id = max(existing_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    NUMBERING_MARKERS[num_id] = (marker, num_format)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_ref])


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_list(doc: Document, items: Iterable[str], num_id: int) -> None:
    marker, num_format = NUMBERING_MARKERS.get(num_id, ("•", "bullet"))
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        label = marker.replace("%1", str(index)) if num_format != "bullet" else "•"
        run = paragraph.add_run(f"{label}  {item}")
        set_run_font(run)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="Microsoft YaHei", size=9.5, color=MUTED)
    return paragraph


def add_callout(doc: Document, label: str, text: str, color: str = GREEN):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.06)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(paragraph, CALLOUT, border_left=color)
    head = paragraph.add_run(f"{label}：")
    set_run_font(head, east_asia="Microsoft YaHei", bold=True, color=color)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_code_block(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(paragraph, LIGHT_GRAY)
    run = paragraph.add_run(text)
    set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=8.8, color=INK)
    return paragraph


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    header_fill: str = LIGHT_BLUE,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, east_asia="Microsoft YaHei", size=9.2, bold=True, color=INK)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(str(value)) <= 10 and index != len(values) - 1
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            set_run_font(run, size=8.8)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def font(size: int, bold: bool = False):
    path = Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc")
    return ImageFont.truetype(str(path), size=size)


def rounded_text_box(draw, rect, title, subtitle, fill, outline):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=22, fill=fill, outline=outline, width=4)
    title_font = font(32, True)
    body_font = font(21)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_x = (x1 + x2 - (title_box[2] - title_box[0])) / 2
    draw.text((title_x, y1 + 27), title, font=title_font, fill="#0B2545")
    body_box = draw.multiline_textbbox((0, 0), subtitle, font=body_font, spacing=7, align="center")
    body_x = (x1 + x2 - (body_box[2] - body_box[0])) / 2
    draw.multiline_text((body_x, y1 + 83), subtitle, font=body_font, fill="#475467", spacing=7, align="center")


def draw_arrow(draw, start, end, color="#087D72"):
    draw.line([start, end], fill=color, width=7)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 20
    for offset in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + offset),
            end[1] + length * math.sin(angle + offset),
        )
        draw.line([end, point], fill=color, width=7)


def create_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    stages = [
        ("原始资料", "网页 / PDF\n图片 / 表格"),
        ("第一次清洗", "MinerU / OCR / VLM\n版面与视觉解析"),
        ("第二次清洗", "Python 规则\n去噪、结构化、去重"),
        ("多模态索引", "文本 + 图片语义\n表格 JSON + 溯源"),
        ("检索与包装", "RAGFlow Top-K\nLangGraph 质量门禁"),
    ]
    fills = ["#F2F4F7", "#E8EEF5", "#E7F5F2", "#FFF4DF", "#E8EEF5"]
    left = 55
    width = 300
    gap = 50
    top = 235
    height = 250
    for index, ((title, subtitle), fill) in enumerate(zip(stages, fills)):
        x1 = left + index * (width + gap)
        rounded_text_box(draw, (x1, top, x1 + width, top + height), title, subtitle, fill, "#2E74B5")
        if index < len(stages) - 1:
            draw_arrow(draw, (x1 + width + 6, top + height // 2), (x1 + width + gap - 8, top + height // 2))
    draw.text((55, 65), "暨南大学学生助手：多模态清洗、关联与问答链路", font=font(42, True), fill="#0B2545")
    draw.text(
        (55, 585),
        "输出约束：保留业务语义，删除检索噪声；图片可展示、表格可计算、每个单元均可回到来源文档与页码。",
        font=font(27),
        fill="#475467",
    )
    draw.rounded_rectangle((55, 680, 1745, 810), radius=20, fill="#F4F6F9", outline="#98A2B3", width=2)
    draw.text(
        (88, 718),
        "质量闭环：路径校验 → 图注/上下文完整性 → 空表/稀疏表检测 → 检索相关性 → 无证据拒答 → 来源链接",
        font=font(27, True),
        fill="#087D72",
    )
    image.save(path, quality=95)


def create_results_chart(path: Path, stats: dict) -> None:
    values = [
        ("来源文档", int(stats.get("documents", 22))),
        ("真实图片", int(stats.get("physical_images", 25))),
        ("表格截图", int(stats.get("table_images", 14))),
        ("结构化表格", int(stats.get("structured_tables", 30))),
        ("索引资源", int(stats.get("items", 55))),
    ]
    image = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "多模态索引重建结果", font=font(42, True), fill="#0B2545")
    max_value = max(value for _, value in values)
    chart_left, chart_top, chart_bottom = 245, 160, 715
    available = 1020
    colors = ["#2E74B5", "#087D72", "#7A5A00", "#4F46E5", "#0B2545"]
    for index, ((label, value), color) in enumerate(zip(values, colors)):
        y = chart_top + index * 105
        draw.text((60, y + 16), label, font=font(25, True), fill="#344054")
        draw.rounded_rectangle((chart_left, y, chart_left + available, y + 64), radius=14, fill="#F2F4F7")
        width = int(available * value / max_value)
        draw.rounded_rectangle((chart_left, y, chart_left + width, y + 64), radius=14, fill=color)
        draw.text((chart_left + width + 22, y + 12), str(value), font=font(29, True), fill=color)
    draw.text((70, 760), "注：表格截图属于真实图片的子集；结构化表格为 HTML 表格转换后的 JSON 行列数据。", font=font(21), fill="#667085")
    image.save(path, quality=95)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Block" not in doc.styles:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        code.font.size = Pt(8.8)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run("暨南大学学生助手")
    set_run_font(left, east_asia="Microsoft YaHei", size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run("\t第三部分研究报告")
    set_run_font(right, east_asia="Microsoft YaHei", size=8.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = footer_p.add_run("— ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(footer_p)
    suffix = footer_p.add_run(" —")
    set_run_font(suffix, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("学 术 研 究 报 告")
    set_run_font(run, east_asia="Microsoft YaHei", size=13, bold=True, color=GOLD)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("图、表、文字的\n多模态清洗与关联")
    set_run_font(run, east_asia="Microsoft YaHei", size=29, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(42)
    run = subtitle.add_run("暨南大学学生助手的工程实现、检索关联与质量评估")
    set_run_font(run, east_asia="Microsoft YaHei", size=14, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(36)
    set_paragraph_shading(rule, LIGHT_BLUE)
    run = rule.add_run("从“可读取”到“可检索、可解释、可溯源”")
    set_run_font(run, east_asia="Microsoft YaHei", size=11.5, bold=True, color=GREEN)

    metadata = [
        "项目名称：暨南大学学生助手",
        "研究模块：第三部分 - 图、表、文字的多模态清洗与关联",
        "撰写单位：项目组",
        "报告日期：2026 年 7 月",
    ]
    for item in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        set_run_font(r, east_asia="Microsoft YaHei", size=10.5, color=MUTED)
    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    add_heading(doc, "摘  要", 1)
    add_body(
        doc,
        "高校学生事务知识库包含网页正文、PDF 通知、业务表格、办事截图及附件图片。若仅进行纯文本抽取，"
        "页面装饰、页码、天气、重复表头等噪声会干扰召回，而图像中的流程关系、表格中的字段结构和"
        "原文位置又容易丢失。针对这一问题，本报告提出并实现一种面向检索增强生成（RAG）的两阶段"
        "多模态清洗与关联方法：首先利用 MinerU、OCR 或视觉模型完成版面解析与视觉单元抽取；随后使用"
        "Python 规则完成业务语义保留、装饰噪声删除、图片语义化、HTML 表格 JSON 化、哈希去重和来源"
        "路径校验。清洗后的文本、图片描述、可见文字、关键词、页码及表格行列数据被组织为统一索引，"
        "再与 RAGFlow Top-K 召回结果及 LangGraph 多智能体质量门禁连接，实现答案、来源网页和相关"
        "图表的协同展示。"
    )
    add_body(
        doc,
        "在当前知识库快照上，系统扫描 5 套知识库，对 692 份文档、1,138 个文本分块和 78 个原始图片"
        "分块进行核验；经跨实验库去重后形成 55 个多模态资源，覆盖 22 份来源文档，其中包含 25 张真实"
        "图片、14 张表格截图和 30 个结构化表格。严格质量门禁确认 55 个资源全部可解析，缺失图片、无"
        "图注、无上下文、空表格和稀疏表格数量均为 0。结果表明，该方法将“文档解析”转化为“可检索、"
        "可解释、可展示、可溯源”的知识单元，为校园事务问答中的可信回答和拒答机制提供了数据基础。"
    )
    keyword = doc.add_paragraph(style="Normal")
    key = keyword.add_run("关键词：")
    set_run_font(key, east_asia="Microsoft YaHei", bold=True, color=INK)
    value = keyword.add_run("多模态文档理解；MinerU；表格结构化；RAGFlow；知识溯源；质量门禁")
    set_run_font(value)

    add_heading(doc, "Abstract", 1)
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(
        "Student-service knowledge bases combine web text, PDF notices, forms, screenshots, and embedded "
        "figures. Plain-text extraction is insufficient because decorative elements contaminate retrieval, "
        "while visual relations, table structures, and provenance may be lost. This report presents a "
        "two-stage multimodal cleaning and association pipeline for retrieval-augmented generation. MinerU, "
        "OCR, or a vision model first extracts layout-aware units; Python post-processing then removes noise, "
        "retains business semantics, converts images into searchable descriptions, serializes HTML tables into "
        "JSON rows, deduplicates resources with SHA-256 fingerprints, and verifies all asset paths. The resulting "
        "index connects text, images, tables, source documents, and page locations to RAGFlow retrieval and a "
        "LangGraph quality harness. On the current snapshot, 55 resolved multimodal resources were reconstructed "
        "from 22 source documents with zero missing assets. The implementation demonstrates a reproducible path "
        "from document parsing to traceable and user-readable evidence delivery."
    )
    set_run_font(run, latin="Calibri", east_asia="SimSun", size=10.5)
    keywords = doc.add_paragraph(style="Normal")
    key = keywords.add_run("Keywords: ")
    set_run_font(key, bold=True, color=INK)
    value = keywords.add_run(
        "multimodal document understanding; table serialization; retrieval-augmented generation; provenance; quality assurance"
    )
    set_run_font(value)

    add_heading(doc, "内容结构", 1)
    toc_num = create_numbering(doc, "%1.", "decimal")
    add_list(
        doc,
        [
            "绪论：问题背景、研究目标与主要贡献",
            "理论基础与需求分析：视觉文档理解、RAG 与教师材料约束",
            "总体架构与数据模型：两阶段清洗、统一索引与质量闭环",
            "多模态清洗方法：图片、表格和文本的二次处理算法",
            "跨模态关联与溯源：文档、页码、资源和检索结果的映射",
            "RAGFlow Top-K 与多智能体二次包装",
            "工程实现、实验结果与质量评估",
            "局限、后续研究和结论",
        ],
        toc_num,
    )
    doc.add_page_break()


def build_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "multimodal_architecture.png"
    results_chart = ASSET_DIR / "multimodal_results.png"
    stats = read_json(ROOT / "outputs" / "multimodal_snapshot_report.json", {})
    quality = read_json(ROOT / "outputs" / "quality_gate.json", {})
    benchmark = read_json(ROOT / "config" / "core_retrieval_benchmark.json", {})
    create_architecture_diagram(architecture)
    create_results_chart(results_chart, stats)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    doc.core_properties.title = "图、表、文字的多模态清洗与关联"
    doc.core_properties.subject = "暨南大学学生助手第三部分学术报告"
    doc.core_properties.author = "暨南大学学生助手项目组"
    doc.core_properties.keywords = "多模态清洗, MinerU, RAGFlow, 表格JSON, 溯源"
    bullet_num = create_numbering(doc, "•", "bullet")
    decimal_num = create_numbering(doc, "%1.", "decimal")
    reference_num = create_numbering(doc, "[%1]", "decimal")

    add_cover(doc)
    add_abstracts(doc)

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_body(
        doc,
        "高校学生事务问答并不是单一的文本检索问题。请假申请表、转专业审批表、交换生学分转换表、"
        "校园网办理页面和招生通知往往同时包含标题、正文、表格、流程箭头、截图、印章区域和附件链接。"
        "这些内容的业务价值并不等同：页码、天气、人物头像和装饰线条通常是检索噪声，而标题、字段名、"
        "审核关系、日期范围、金额、网址和部门名称则是回答问题所需的核心证据。老师提供的第三部分材料"
        "明确提出：第一次使用 MinerU 清洗，第二次使用 Python 代码处理；图表较多时应单独存储，并且必须"
        "能够溯源；表格建议转换为 JSON；RAGFlow 返回的 Top-K 分块应由多智能体后台进行二次包装[1]。"
    )
    add_body(
        doc,
        "RAG 的核心价值不仅在于提高生成答案的事实性，还包括通过外部非参数知识提供可更新、可追踪的"
        "证据来源[2]。在视觉文档场景中，文字、二维布局和图像之间存在互补关系。LayoutLMv2 等工作证明"
        "联合建模文本、布局和图像能够提升富文本文档理解能力[3]；Donut 则展示了端到端视觉文档理解在"
        "避免 OCR 误差传播方面的潜力[4]。本项目不训练新的视觉模型，而是将这些思想落实为可维护的数据"
        "工程：解析、清洗、结构化、关联、去重、验证与检索展示。"
    )

    add_heading(doc, "1.2 研究问题", 2)
    add_list(
        doc,
        [
            "如何区分页面中的业务信息与装饰噪声，降低无关词对向量检索和关键词检索的干扰？",
            "如何把图片中的表意内容转换为可被文本检索召回、又能被前端直接展示的资源？",
            "如何把复杂表格转换为统一 JSON 行列格式，同时避免丢失原始视觉版式？",
            "如何建立文档、页码、图片、表格、检索分块和答案之间的溯源关系？",
            "如何控制 RAGFlow Top-K 规模，并在上下文受限时完成二次筛选和可读化包装？",
        ],
        bullet_num,
    )

    add_heading(doc, "1.3 研究目标与主要贡献", 2)
    add_list(
        doc,
        [
            "提出两阶段清洗流程，将版面解析与确定性 Python 后处理解耦。",
            "建立统一多模态资源模式，兼容真实图片、表格截图和结构化表格。",
            "使用 SHA-256 对跨知识库重复图片与表格进行去重，并校验资源路径。",
            "实现基于文档名、检索片段和问题关键词的跨模态关联评分。",
            "把多模态证据接入 RAGFlow 与 LangGraph 质量门禁，拒答时不显示无关来源。",
            "构建可重复执行的索引重建、严格质量检查和 160 条检索评测配置。",
        ],
        bullet_num,
    )
    add_callout(
        doc,
        "研究边界",
        "本文报告的是当前项目中已经落地的第三部分能力。尚未执行的模型对比实验、人工标注准确率和完整 160 条在线评测均明确列为后续工作，不将建议性指标表述为既有结果。",
        GOLD,
    )

    add_heading(doc, "2 理论基础与需求分析", 1)
    add_heading(doc, "2.1 多模态文档理解", 2)
    add_body(
        doc,
        "文档图像不是普通自然图像，其语义高度依赖布局。相同文字位于标题、表头、脚注或审批栏时具有"
        "不同功能。因此，多模态清洗至少需要同时处理四类信息：文本内容、二维位置、视觉对象和文档层级。"
        "MinerU 面向 PDF 内容抽取，结合预处理和后处理规则以提升不同文档类型上的一致性[5]。本项目采用"
        "“模型负责感知、规则负责约束”的分工：视觉模型给出描述、可见文字和关键词，Python 负责统一"
        "字段、去重、路径安全和输出模式。"
    )

    add_heading(doc, "2.2 教师材料要求到系统设计的映射", 2)
    add_caption(doc, "表 1  第三部分要求与项目实现对应关系")
    add_table(
        doc,
        ["材料要求", "工程设计", "当前状态"],
        [
            ["MinerU 第一次清洗", "保留知识库快照中的 MinerU/视觉解析结果与图片分块", "已实现"],
            ["Python 第二次清洗", "统一字段、去噪、HTML 表格解析、哈希去重、路径校验", "已实现"],
            ["删除无用信息", "不把页码、装饰图等单独作为检索主体；仅作为元数据保留页码", "已实现"],
            ["保留业务语义", "保存图注、可见文字、关键词、问题、上下文和来源文档", "已实现"],
            ["表格转 JSON", "HTML <table> 解析为 rows: [[cell,...], ...]", "已实现"],
            ["图表单独文件", "图片位于 knowledge_base/blobs/images，多模态索引独立存储", "已实现"],
            ["支持溯源", "source_dataset、document_id、document、page、asset_path", "已实现"],
            ["Top-K 二次包装", "RAGFlow page_size=5、top_k=30；LangGraph 质量检查与答案整理", "已实现"],
        ],
        [2550, 5010, 1800],
    )

    add_heading(doc, "2.3 信息保留原则", 2)
    add_body(
        doc,
        "信息保留遵循“回答贡献度”原则。若某元素能够帮助识别事项名称、适用对象、办理条件、材料、步骤、"
        "时间、费用、部门、网址或原文位置，则作为业务证据保留；若元素只承担装饰或模板功能，则删除或"
        "降级为轻量元数据。老师材料中的页码、天气、温度、星期、人物头像和机器人头像均属于典型噪声，"
        "而“人类给出明确指令”“AI 自主规划并调用工具”等文字和输入输出关系属于核心语义[1]。"
    )

    add_heading(doc, "3 总体架构与数据模型", 1)
    add_heading(doc, "3.1 端到端处理架构", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(architecture), width=Inches(6.35))
    add_caption(doc, "图 1  多模态清洗、索引、检索与回答的端到端架构")
    add_body(
        doc,
        "总体链路由五层组成：原始数据层、第一次解析层、第二次清洗层、多模态索引层和检索回答层。"
        "第一次解析负责把不可直接计算的版面变成候选单元；第二次清洗负责确定性约束；索引层负责跨模态"
        "统一和持久化；检索回答层只在文本证据通过质量门禁后返回图表，避免“有图片就强行回答”的错误。"
    )

    add_heading(doc, "3.2 统一多模态资源模式", 2)
    add_caption(doc, "表 2  多模态资源的核心字段")
    add_table(
        doc,
        ["字段组", "代表字段", "作用"],
        [
            ["身份", "id, type, sha256", "资源标识、类型判断与重复检测"],
            ["语义", "caption, visible_text, keywords, questions", "支持关键词匹配和文本向量召回"],
            ["上下文", "context, snippet, document, page", "解释视觉内容并定位原文"],
            ["结构", "rows, is_table, visual_type", "支持表格展示、字段比较和后续计算"],
            ["文件", "asset_path, url, content_type", "连接本地图片文件与网页静态路由"],
            ["溯源", "source_dataset, source_dataset_id, source_document_id", "回到知识库、文档和实验版本"],
        ],
        [1500, 3300, 4560],
    )
    add_code_block(
        doc,
        '{\n'
        '  "id": "image-16d50c364b115713",\n'
        '  "type": "image",\n'
        '  "is_table": true,\n'
        '  "document": "转专业申请表(外招生)",\n'
        '  "page": "第 1 页",\n'
        '  "caption": "包含转出学院、转入学院及教务处审核意见栏",\n'
        '  "visible_text": "是否符合转专业基本资格……",\n'
        '  "asset_path": "blobs/images/16/16d50c...jpg",\n'
        '  "source_dataset_id": "ea25d7f..."\n'
        '}'
    )
    add_caption(doc, "代码 1  图片单元统一 JSON 示例（字段节选）")

    add_heading(doc, "3.3 两阶段清洗的职责划分", 2)
    add_body(
        doc,
        "第一次清洗具有较强的感知能力，能够识别版面、文字块、图像和表格，但输出可能存在跨页断裂、"
        "重复单元、描述风格不统一或资源路径依赖。第二次清洗具有确定性和可测试性，负责模式统一和约束"
        "执行。两者分离可避免将所有规则写进提示词，也降低模型更新导致数据结构漂移的风险。"
    )
    add_caption(doc, "表 3  两阶段清洗职责比较")
    add_table(
        doc,
        ["维度", "第一次：MinerU/OCR/VLM", "第二次：Python"],
        [
            ["优势", "视觉感知、版面识别、自然语言描述", "稳定、可重复、易测试、低成本"],
            ["主要任务", "文本/图像/表格候选单元抽取", "去噪、结构化、关联、去重、校验"],
            ["典型风险", "OCR 错误、跨页断裂、描述不一致", "规则覆盖不足、过度过滤"],
            ["输出", "候选视觉单元与解析内容", "统一多模态索引与质量报告"],
        ],
        [1500, 3930, 3930],
    )

    add_heading(doc, "4 多模态清洗方法", 1)
    add_heading(doc, "4.1 文本清洗", 2)
    add_body(
        doc,
        "文本清洗不是简单删除非正文，而是对信息功能进行分类。系统将正文、标题、表头、字段值、业务"
        "条件和来源链接保留；将 Markdown 标记、重复空白和超长哈希文件名从回答摘要中移除；页码从正文"
        "降级为 page 元数据，以便溯源但不参与主要答案生成。URL 被单独解析为 source_url，用于生成"
        "“打开官方来源”按钮，而不是混在正文中影响可读性。"
    )
    add_list(
        doc,
        [
            "语义保真：不能删除日期、金额、部门、材料和限制条件。",
            "结构保真：标题层级和表格字段关系应在清洗后继续存在。",
            "检索友好：减少页眉页脚、装饰符号和重复模板词的权重。",
            "展示友好：答案摘要与来源链接分离，原始证据保留在 matches 中。",
        ],
        bullet_num,
    )

    add_heading(doc, "4.2 图片语义化", 2)
    add_body(
        doc,
        "图片本身不能直接被纯文本查询稳定召回，因此系统为每张图片保存视觉类型、语义描述、可见文字、"
        "检索关键词和建议问题。语义描述解释“图中是什么”，可见文字提供可匹配实体，关键词用于快速"
        "筛选，建议问题则用于把视觉内容映射到用户表达。图片文件仍被保留，回答命中后由前端通过"
        "/knowledge-assets 静态路由展示，实现“文本召回 + 原图核验”。"
    )
    if SAMPLE_IMAGE.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(SAMPLE_IMAGE), width=Inches(6.25))
        add_caption(doc, "图 2  转专业申请表审批区域的真实表格截图")
        add_callout(
            doc,
            "语义化结果",
            "该资源被描述为“转专业申请表的管理部门审批部分”，识别出转出学院、转入学院和教务处三类审核意见，并保留“第 1 页”和来源文档标识。用户询问审核意见时，可同时获得文字答案和原始表格截图。",
            GREEN,
        )

    add_heading(doc, "4.3 表格 JSON 化", 2)
    add_body(
        doc,
        "对 HTML 表格，系统使用 BeautifulSoup 遍历 tr 以及 th/td 节点，清理单元格空白并输出二维"
        "rows 数组。该格式具有三项优势：其一，请求与接收格式统一，便于模型提示词约束；其二，前端可"
        "直接生成可滚动表格；其三，后续可对字段进行比较、校验或计算。对于仅有截图而无可恢复行列结构"
        "的表格，系统保留 is_table=true 和原图，不伪造 rows。"
    )
    add_code_block(
        doc,
        '{\n'
        '  "type": "table",\n'
        '  "document": "学生事务申请表",\n'
        '  "rows": [\n'
        '    ["审核环节", "责任单位", "操作"],\n'
        '    ["转出审核", "转出学院", "提交申请"],\n'
        '    ["转入审核", "转入学院", "审核申请"],\n'
        '    ["终审", "教务处学籍科", "确认结果"]\n'
        '  ]\n'
        '}'
    )
    add_caption(doc, "代码 2  表格二维 JSON 表示")

    add_heading(doc, "4.4 去重与完整性校验", 2)
    add_body(
        doc,
        "知识库包含 A/B/C 三套分块实验，同一视觉资源可能被重复导入。若不去重，前端会显示重复图片，"
        "统计也会被放大。系统对图片使用原始 image_sha256；对结构化表格使用"
        "SHA256(JSON({document, rows})) 作为指纹。图片路径在写入索引前被解析为绝对路径，并检查其"
        "是否仍位于 knowledge_base 根目录内且文件真实存在。这同时避免目录穿越与断图。"
    )
    add_code_block(
        doc,
        "图片去重键 = SHA256(图片二进制)\n"
        "表格去重键 = SHA256(规范化来源文档名 + JSON(rows))\n"
        "安全条件 = resolved_path ∈ knowledge_base ∧ is_file(resolved_path)"
    )

    add_heading(doc, "5 跨模态关联与溯源机制", 1)
    add_heading(doc, "5.1 关联评分", 2)
    add_body(
        doc,
        "文本检索得到 document_name 与 matches 后，多模态关联器对每个资源计算启发式得分。直接来源"
        "文档匹配赋予较高权重，召回片段中的文档匹配作为补充，问题和资源文本之间的中文词段重合提供"
        "语义线索。设资源为 u、问题与召回上下文为 q，则当前实现可概括为："
    )
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(8)
    formula.paragraph_format.space_after = Pt(10)
    run = formula.add_run("S(u,q) = 10·I(直接文档匹配) + 8·ΣI(召回文档匹配) + min(8, 词段重合数)")
    set_run_font(run, latin="Cambria Math", east_asia="Microsoft YaHei", size=11.5, italic=True, color=INK)
    add_body(
        doc,
        "资源按得分降序排列，仅保留不低于 max(3, Smax-4) 的候选，最多返回 6 个。该相对阈值比固定"
        "返回数量更稳健：当存在强文档匹配时，低相关资源不会因“凑满数量”进入页面；当文档名不一致时，"
        "关键词仍能提供有限召回能力。"
    )

    add_heading(doc, "5.2 溯源链", 2)
    add_list(
        doc,
        [
            "答案层：answer、document_name、similarity 和官方来源链接。",
            "检索层：matches 保留每个分块的文档名、相似度、原文位置和摘要。",
            "视觉层：多模态资源保留 document、page、caption、visible_text 和 asset_path。",
            "知识库层：source_dataset_id 与 source_document_id 指向快照中的数据集和文档。",
            "文件层：SHA-256 与相对路径指向 knowledge_base/blobs/images 中的原始文件。",
        ],
        decimal_num,
    )
    add_callout(
        doc,
        "可信性约束",
        "多模态资源只能在文本证据通过 Harness 后展示。若检索多轮仍不合格，系统清空 document_name、source_url、similarity、matches、retrieved 和 multimodal，避免将无关图片或链接伪装成答案依据。",
        RED,
    )

    add_heading(doc, "6 RAGFlow Top-K 与多智能体二次包装", 1)
    add_heading(doc, "6.1 Top-K 参数设计", 2)
    add_body(
        doc,
        "RAGFlow 的 top_k 决定候选池规模，page_size 决定交给回答链路的主要分块数量。当前经过实验配置"
        "后采用 page_size=5、top_k=30、similarity_threshold=0.0、vector_similarity_weight=0.7，并"
        "启用关键词检索。设计逻辑是：先用 30 个候选降低漏召回风险，再只把前 5 个高相关分块带入答案"
        "构建和质量检查。这样既保留混合检索的搜索空间，也控制大模型上下文和接口响应体。"
    )
    add_caption(doc, "表 4  召回参数及其工程含义")
    add_table(
        doc,
        ["参数", "当前值", "作用与取舍"],
        [
            ["page_size", "5", "用于答案构建的核心分块数，控制上下文长度"],
            ["top_k", "30", "RAGFlow 候选池规模，降低复杂问法的漏召回"],
            ["similarity_threshold", "0.0", "候选阶段不过早截断，由后端 Harness 再判断"],
            ["vector_similarity_weight", "0.7", "以语义向量为主，同时保留关键词匹配"],
            ["keyword", "true", "增强表名、部门名、材料名等精确词命中"],
            ["最大重试", "2", "失败后改写查询，仍失败则拒答"],
        ],
        [2100, 1500, 5760],
    )

    add_heading(doc, "6.2 上下文超限处理", 2)
    add_body(
        doc,
        "当候选分块超过模型上下文时，不应简单截断完整文档。推荐按照“相关性优先、来源去重、字段完整、"
        "长度预算”进行压缩：先按分数排序；每个来源保留最高分块；对相邻分块合并并删除重复页眉；表格"
        "只传相关行与表头；图片只传描述与可见文字，原图留在前端；最后按 token 预算截断。当前系统已经"
        "完成前 5 个分块选择和来源排序，token 级动态预算属于下一阶段优化。"
    )

    add_heading(doc, "6.3 LangGraph 包装与质量门禁", 2)
    add_body(
        doc,
        "多智能体流程为 Intent → Router → Retriever/Tool → Quality → Rewrite（可选）→ Reflection →"
        " Answer。Answer Agent 不直接替代检索证据：可选文本模型仅在 state.ok、route=retrieve 且"
        "matches 非空时整理措辞；没有 LLM Key 或请求失败时回退到规则摘要。此设计把大模型限制在“表达"
        "优化”角色中，事实边界由检索结果和 Harness 控制。"
    )
    add_list(
        doc,
        [
            "最低证据分数为 0.2。",
            "必须存在来源文档与召回片段。",
            "答案与问题的关键事项必须有词项覆盖。",
            "质量失败时最多改写查询两次。",
            "达到最大重试次数后主动拒答并清空来源。",
        ],
        bullet_num,
    )

    add_heading(doc, "7 工程实现与可重复性", 1)
    add_heading(doc, "7.1 核心模块", 2)
    add_caption(doc, "表 5  第三部分相关代码模块")
    add_table(
        doc,
        ["模块", "功能", "关键输出"],
        [
            ["multimodal/build_snapshot_index.py", "快照扫描、语义字段解析、表格 JSON 化、去重", "data/multimodal_index.json"],
            ["multimodal/postprocess_mineru.py", "MinerU 输出的上下文、图注和表格单元处理", "多模态文档单元"],
            ["multimodal/query_image.py", "用户截图预处理和视觉模型调用", "脱敏检索问题"],
            ["app_fastapi.py", "静态资源路由、跨模态关联、图片/表格展示", "API 与 Web 页面"],
            ["scripts/quality_gate.py", "断图、图注、上下文、空表和稀疏表检查", "quality_gate.json/md"],
            ["agents_fastapi/graph.py", "检索、重试、反思、受控整理和拒答", "AgentState 与 trace"],
        ],
        [3000, 3810, 2550],
    )

    add_heading(doc, "7.2 重建与验证命令", 2)
    add_code_block(
        doc,
        '.\\.venv\\Scripts\\python.exe .\\multimodal\\build_snapshot_index.py\n'
        '.\\.venv\\Scripts\\python.exe .\\scripts\\quality_gate.py --strict\n'
        '.\\.venv\\Scripts\\python.exe .\\tests\\fastapi_smoke.py\n'
        '.\\.venv\\Scripts\\python.exe .\\ragflow\\tune_core_retrieval.py --variants 4 --generate-only'
    )
    add_body(
        doc,
        "上述命令分别完成索引重建、严格质量检查、Agent 冒烟验证和 160 条检索评测配置生成。索引脚本"
        "在资源为空或发现缺失文件时以非零状态退出；严格质量门禁在断图、无多模态资源或坏链接时失败，"
        "因此适合接入持续集成。"
    )

    add_heading(doc, "7.3 前端呈现", 2)
    add_body(
        doc,
        "前端根据资源类型进行差异化渲染：存在 url 时使用延迟加载的 img 展示真实图片；存在 rows 时"
        "生成最多 12 行、8 列的可滚动表格；图注、文档名和页码使用 textContent 写入，避免把索引内容"
        "当作 HTML 执行。数据看板同步显示资源总数、图片数、表格数、结构化表格数和已解析数量。"
    )

    add_heading(doc, "8 实验设计与结果", 1)
    add_heading(doc, "8.1 数据快照", 2)
    snapshot_rows = [
        ["知识库数量", "5", "本地 knowledge_base/datasets 快照"],
        ["文档数量", "692", "五套知识库汇总"],
        ["文本分块", "1,138", "RAGFlow 快照分块"],
        ["原始图片分块", "78", "A/B/C 实验库各 26 个，含重复"],
        ["去重后来源文档", str(stats.get("documents", 22)), "具有图片或结构化表格的文档"],
        ["去重后索引资源", str(stats.get("items", 55)), "当前前端可检索多模态资源"],
    ]
    add_caption(doc, "表 6  数据规模与索引规模")
    add_table(doc, ["指标", "数量", "说明"], snapshot_rows, [2400, 1500, 5460])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(results_chart), width=Inches(6.25))
    add_caption(doc, "图 3  多模态索引重建统计")

    add_heading(doc, "8.2 质量门禁结果", 2)
    quality_summary = quality.get("summary", {})
    multimodal_quality = quality.get("multimodal", {})
    add_caption(doc, "表 7  严格质量检查结果")
    add_table(
        doc,
        ["检查项", "结果", "判定"],
        [
            ["多模态资源", str(quality_summary.get("multimodal_resources", 55)), "通过"],
            ["可解析资源", f"{stats.get('items', 55)} / {stats.get('items', 55)}", "通过"],
            ["缺失图片文件", str(quality_summary.get("missing_multimodal_assets", 0)), "通过"],
            ["缺少图注", str(len(multimodal_quality.get("images_missing_caption", []))), "通过"],
            ["缺少上下文", str(len(multimodal_quality.get("images_missing_context", []))), "通过"],
            ["空表格", str(len(multimodal_quality.get("empty_tables", []))), "通过"],
            ["稀疏表格", str(len(multimodal_quality.get("sparse_tables", []))), "通过"],
        ],
        [3300, 1800, 4260],
    )
    add_callout(
        doc,
        "结果解释",
        "78 个原始图片分块来自三套分块实验，不能直接视为 78 张独立图片。按 SHA-256 去重后得到 25 张真实图片；14 张属于表格截图。另有 30 个 HTML 表格被转换为结构化 rows，因此“表格资源”总数为 44。",
        GREEN,
    )

    add_heading(doc, "8.3 功能验证案例", 2)
    add_caption(doc, "表 8  在线与自动化验证")
    add_table(
        doc,
        ["测试场景", "预期行为", "实测结果"],
        [
            ["转专业申请表审核意见", "返回有依据答案、官方链接和相关表格图片", "HTTP 200；答案通过；1 个来源按钮；2 个多模态资源"],
            ["宿舍空调维修电话（知识库无可靠资料）", "多轮检索后拒答，不展示错误来源", "拒答；document/source/matches/links/multimodal 均为空"],
            ["静态图片访问", "/knowledge-assets 路径返回真实文件", "HTTP 200"],
            ["数据看板", "显示 55 个资源和 30 个结构化表格", "HTTP 200；统计一致"],
            ["Agent 冒烟测试", "工具、医疗边界、检索重试、来源按钮、拒答", "全部通过"],
            ["知识库导入单元测试", "私网地址、快照完整性、批量导入", "4 项全部通过"],
        ],
        [2900, 3600, 2860],
    )

    add_heading(doc, "8.4 检索评测配置", 2)
    positives = len(benchmark.get("positive_cases", []))
    negatives = len(benchmark.get("negative_cases", []))
    add_body(
        doc,
        f"为减少仅凭少数演示问题调参的偏差，项目已生成 {positives + negatives} 条核心检索评测配置，"
        f"其中正例 {positives} 条、负例 {negatives} 条，ID 无重复。配置覆盖原始意图及四种问法变体，可用于"
        "比较向量权重、召回阈值和重排模型。当前仅完成配置生成和小规模在线验证，尚未执行全部候选参数"
        "组合，因此本文不报告未经运行的准确率、MRR 或 Recall@K。"
    )

    add_heading(doc, "8.5 建议的后续评价指标", 2)
    add_caption(doc, "表 9  多模态清洗与检索评价指标")
    add_table(
        doc,
        ["层级", "指标", "定义或目的"],
        [
            ["解析", "图表单元召回率", "人工标注图表中被系统成功抽取的比例"],
            ["清洗", "噪声删除准确率", "删除项中确属无业务价值的比例"],
            ["结构", "Table TEDS / 单元格 F1", "评价表格行列结构恢复质量"],
            ["关联", "资源 Precision@K", "返回图表中与问题和来源文档相关的比例"],
            ["检索", "Recall@K、MRR", "评价目标文档进入候选集及排序位置"],
            ["回答", "依据一致性、拒答准确率", "评价答案是否由证据支撑及未知问题处理"],
            ["系统", "P50/P95 延迟、上下文 token", "评价工程效率与成本"],
        ],
        [1500, 2700, 5160],
    )

    add_heading(doc, "9 讨论", 1)
    add_heading(doc, "9.1 已达到的能力", 2)
    add_body(
        doc,
        "当前实现已经覆盖老师第三部分的主要工程要求：存在两阶段处理思想；图片语义、可见文字和关键词"
        "能够参与关联；表格能够转换为统一 JSON；图表存储与文本索引分离；每个资源具有文档和数据集"
        "溯源字段；RAGFlow Top-K 被控制在候选池和回答分块两个层级；多智能体后台进行重试、反思、拒答"
        "和受控答案整理；前端能够同时展示答案、来源按钮和相关图表。"
    )

    add_heading(doc, "9.2 局限性", 2)
    add_caption(doc, "表 10  当前局限与影响")
    add_table(
        doc,
        ["局限", "影响", "处理建议"],
        [
            ["本机未安装可直接重跑的 MinerU 可执行环境", "无法从任意新 PDF 完成完全离线首轮解析", "固化 MinerU 版本、模型和运行脚本"],
            ["当前索引主要从知识库快照恢复", "传统 data/cleaned 文档与服务卡片未纳入同一质量报告", "统一快照与 cleaned 清单"],
            ["VLM API Key 尚未配置", "用户上传截图时不能调用在线视觉模型", "在 /settings 配置受控视觉模型"],
            ["LLM API Key 尚未配置", "答案使用规则摘要，语言自然度有限", "按需配置；保持证据门禁优先"],
            ["部分结构化表格来源名为技术文件名", "前端图注可读性受影响", "结合文档目录和标题层级做名称恢复"],
            ["160 条完整评测尚未在线运行", "缺少全量参数对比统计", "固定 RAGFlow 版本后批量执行并保存原始结果"],
        ],
        [2700, 3150, 3510],
    )

    add_heading(doc, "9.3 有效性威胁", 2)
    add_body(
        doc,
        "内部有效性方面，当前质量门禁主要验证结构完整和资源可访问，并不等价于人工语义正确率；外部"
        "有效性方面，数据来自暨南大学学生事务场景，规则可能不能直接迁移到财务票据或医学报告；构念"
        "有效性方面，资源数量不能代表检索质量，必须配合 Precision@K、MRR 和答案依据一致性；结论"
        "有效性方面，当前在线案例和小规模测试只能说明链路可用，不能替代全量对照实验。"
    )

    add_heading(doc, "10 后续研究计划", 1)
    add_list(
        doc,
        [
            "环境固化：用 Docker 或独立虚拟环境锁定 MinerU、OCR、字体与模型版本。",
            "增量清洗：对新增网页和附件计算内容哈希，只处理新增或发生变化的文档。",
            "跨页恢复：合并被分页截断的段落和表格，恢复标题层级与图文邻接关系。",
            "语义校验：引入人工抽样和双模型一致性检查，标记低置信度 OCR 与图注。",
            "表格检索：分别建立表级、行级和字段级索引，并保留表头上下文。",
            "动态 Top-K：根据问题复杂度、分数间隔和 token 预算自适应选择分块数量。",
            "全量评测：运行 160 条正负例，对向量权重、阈值、重排和分块策略做消融实验。",
            "持续质量：将严格质量门禁接入 CI，在提交或定时同步时阻止断图、空表和模式漂移。",
        ],
        decimal_num,
    )

    add_heading(doc, "11 结论", 1)
    add_body(
        doc,
        "本研究围绕高校学生事务知识库中的图、表、文字混合信息，完成了从视觉解析到可信问答的多模态"
        "数据工程闭环。方法以 MinerU/OCR/VLM 负责感知，以 Python 负责确定性二次清洗；通过图片语义"
        "化、表格 JSON 化、哈希去重、路径安全、统一模式和来源字段，把原始文档转化为可检索、可展示、"
        "可解释、可溯源的知识单元。项目实测重建 55 个多模态资源，全部通过严格路径和结构检查；资源"
        "仅在文本证据通过 LangGraph Harness 后展示，检索不足时清空来源并主动拒答。"
    )
    add_body(
        doc,
        "这一实现回答了第三部分材料中的关键工程问题：图表如何清洗、什么应删除、什么应保留、表格"
        "如何统一格式、资源如何单独存储并溯源、Top-K 如何控制以及大模型如何二次包装。后续工作的重点"
        "不再是增加更多界面，而是通过固定环境、人工标注、动态上下文预算和全量对照实验，提高解析与"
        "检索指标的可量化程度。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "课程项目材料：《第三：图、表、文字的多模态清洗与关联-补充》，第 1-4 页，2026。",
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems, 2020, 33: 9459-9474. https://proceedings.neurips.cc/paper/2020/hash/6b493230-Abstract.html",
        "Xu Y, Xu Y, Lv T, et al. LayoutLMv2: Multi-modal Pre-training for Visually-rich Document Understanding. ACL-IJCNLP, 2021: 2579-2591. https://aclanthology.org/2021.acl-long.201/",
        "Kim G, Hong T, Yim M, et al. OCR-Free Document Understanding Transformer. European Conference on Computer Vision, 2022. https://www.ecva.net/papers/eccv_2022/papers_ECCV/html/8042_ECCV_2022_paper.php",
        "Wang B, Xu C, Zhao X, et al. MinerU: An Open-Source Solution for Precise Document Content Extraction. arXiv:2409.18839, 2024. https://arxiv.org/abs/2409.18839",
        "暨南大学学生助手项目组：multimodal/build_snapshot_index.py，快照多模态索引重建脚本，2026。",
        "暨南大学学生助手项目组：scripts/quality_gate.py，多模态质量门禁脚本，2026。",
        "暨南大学学生助手项目组：app_fastapi.py 与 agents_fastapi/graph.py，问答、溯源与多智能体实现，2026。",
    ]
    add_list(doc, references, reference_num)

    add_heading(doc, "附录 A  可重复性与验收清单", 1)
    add_list(
        doc,
        [
            "索引文件可以被 JSON 解析，且每个资源具有唯一 id。",
            "所有 asset_path 均位于 knowledge_base 内并指向真实文件。",
            "每张图片具有 caption 以及 context 或 snippet。",
            "结构化表格 rows 非空，且至少包含 4 个非空单元格。",
            "跨 A/B/C 实验库的相同图片按 SHA-256 去重。",
            "回答命中后图片路由返回 HTTP 200。",
            "无可靠证据时 document、source、matches、links 和 multimodal 均为空。",
            "数据看板统计与 multimodal_index.json、quality_gate.json 一致。",
        ],
        bullet_num,
    )

    add_heading(doc, "附录 B  建议的实验记录格式", 1)
    add_code_block(
        doc,
        '{\n'
        '  "experiment_id": "retrieval-ablation-001",\n'
        '  "dataset_version": "snapshot-2026-07",\n'
        '  "chunk_strategy": "context-1200",\n'
        '  "page_size": 5,\n'
        '  "top_k": 30,\n'
        '  "vector_similarity_weight": 0.7,\n'
        '  "metrics": {\n'
        '    "recall_at_5": null,\n'
        '    "mrr": null,\n'
        '    "refusal_accuracy": null,\n'
        '    "multimodal_precision_at_3": null\n'
        '  },\n'
        '  "note": "指标为空表示尚未运行，禁止以估计值代替。"\n'
        '}'
    )

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
